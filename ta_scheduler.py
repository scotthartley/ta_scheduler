import copy
import csv
import datetime
import json
import os
import random
import re
import socket
import threading
import traceback

from flask import Flask, jsonify, request, send_file

# When frozen by PyInstaller, data files live under sys._MEIPASS.
_BASE_DIR = getattr(__import__("sys"), "_MEIPASS",
                    os.path.dirname(os.path.abspath(__file__)))

app = Flask(__name__, static_folder=os.path.join(_BASE_DIR, "static"))

# None until the user opens or saves a file.
_data_file = None

# Set to the pywebview window once it's created, or None if running headless.
_window = None

EMPTY_DATA = {
    "roles": [],
    "grad_courses": [], "labs": [], "tas": [], "assignments": [],
    "exams": [], "proctor_assignments": [],
    "settings": {},
}


def get_data_file():
    return _data_file


def set_data_file(path):
    global _data_file
    _data_file = path


# ── helpers ──────────────────────────────────────────────────────────────────

def load_data():
    p = get_data_file()
    if not p or not os.path.exists(p):
        return copy.deepcopy(EMPTY_DATA)
    with open(p) as f:
        data = json.load(f)
    # Backward compat: fill missing keys
    for key in EMPTY_DATA:
        if key not in data:
            data[key] = copy.deepcopy(EMPTY_DATA[key])
    return data


def save_data_to_file(data):
    with open(get_data_file(), "w") as f:
        json.dump(data, f, indent=2)


def _file_dialog(dialog_type, directory=None, save_filename="", file_types=()):
    """Show a native file dialog via pywebview and return the chosen path, or None."""
    if _window is None:
        return None
    from webview import FileDialog
    kwargs = dict(directory=directory or _default_dir(), file_types=file_types)
    if dialog_type == FileDialog.SAVE:
        kwargs["save_filename"] = save_filename
    result = _window.create_file_dialog(dialog_type, **kwargs)
    if not result:
        return None
    # SAVE dialog returns a plain string; OPEN/FOLDER returns a tuple
    if isinstance(result, str):
        return result or None
    return result[0] if len(result) > 0 else None


def times_overlap(s1, e1, s2, e2):
    return s1 < e2 and e1 > s2


def _get_meetings(item):
    """Return list of (day, start_min, end_min) for all meetings of an item."""
    if item.get("meetings"):
        return [(m["day"], m["start_min"], m["end_min"]) for m in item["meetings"]]
    day = item.get("day")
    if day is not None:
        return [(day, item.get("start_min", 0), item.get("end_min", 0))]
    return []


def _dc_dates(dc):
    """(start_date, end_date) raw ISO strings for a date_conflicts[] entry,
    falling back to the legacy singular `date` field."""
    start_date = dc.get("start_date") or dc.get("date")
    end_date = dc.get("end_date") or dc.get("date") or start_date
    return start_date, end_date


def _date_conflict_days(dc):
    """Parsed (d0, d1) datetime.date pair with d0 <= d1, or None if unparseable."""
    start_date, end_date = _dc_dates(dc)
    if not start_date or not end_date:
        return None
    try:
        d0 = datetime.date.fromisoformat(start_date)
        d1 = datetime.date.fromisoformat(end_date)
    except ValueError:
        return None
    if d1 < d0:
        d0, d1 = d1, d0
    return d0, d1


def _date_conflict_window_for_day(dc, days, day):
    """(start_min, end_min) covering `day` within the original (unclamped) span
    `days`, or None if `day` falls outside it. The first day uses start_min,
    the last day uses end_min, and days in between are treated as full days."""
    d0, d1 = days
    if day < d0 or day > d1:
        return None
    start_min = dc.get("start_min", 0) if day == d0 else 0
    end_min = dc.get("end_min", 1440) if day == d1 else 1440
    return start_min, end_min


def _expand_date_conflict_weekdays(dc, clamp_start=None, clamp_end=None):
    """Expand a date_conflicts[] entry into [(weekday, start_min, end_min), ...],
    one tuple per calendar day the span touches. `clamp_start`/`clamp_end` (ISO
    strings) narrow the iteration bounds to their intersection with the entry's
    real span — they never change which day counts as first/last for partial-
    time-window purposes, which is always based on the original span."""
    days = _date_conflict_days(dc)
    if days is None:
        return []
    d0, d1 = days
    iter_start, iter_end = d0, d1
    if clamp_start:
        try:
            cs = datetime.date.fromisoformat(clamp_start)
            iter_start = max(iter_start, cs)
        except ValueError:
            pass
    if clamp_end:
        try:
            ce = datetime.date.fromisoformat(clamp_end)
            iter_end = min(iter_end, ce)
        except ValueError:
            pass
    result = []
    day = iter_start
    while day <= iter_end:
        window = _date_conflict_window_for_day(dc, days, day)
        if window is not None:
            result.append((day.weekday(), window[0], window[1]))
        day += datetime.timedelta(days=1)
    return result


def _date_conflict_overlaps(dc, date_obj, start_min, end_min):
    """True if a date_conflicts[] entry overlaps a specific literal date/time
    window (used by the proctoring solver, where exams have exact dates)."""
    days = _date_conflict_days(dc)
    if days is None:
        return False
    d0, d1 = days
    if date_obj < d0 or date_obj > d1:
        return False
    window = _date_conflict_window_for_day(dc, days, date_obj)
    if window is None:
        return False
    return times_overlap(start_min, end_min, window[0], window[1])


def fmt_time(minutes):
    h, m = divmod(int(minutes), 60)
    ampm = "PM" if h >= 12 else "AM"
    h12 = h - 12 if h > 12 else (12 if h == 0 else h)
    return f"{h12}:{m:02d} {ampm}"


DAY_SHORT = ["Mon", "Tue", "Wed", "Thu", "Fri"]
DAY_LONG = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]


def fmt_meetings(item, day_names=None):
    """Every meeting of an item: 'Mon 9:00 AM–10:15 AM, Wed 9:00 AM–10:15 AM'."""
    names = day_names or DAY_SHORT
    parts = [f"{names[d]} {fmt_time(s)}–{fmt_time(e)}"
             for d, s, e in _get_meetings(item) if 0 <= d < len(names)]
    return ", ".join(parts) if parts else "—"


# ── CSV import helpers ────────────────────────────────────────────────────────

def _parse_days(s):
    """'MWF' → [0,2,4], 'TR' → [1,3]"""
    DAY_MAP = {'M': 0, 'T': 1, 'W': 2, 'R': 3, 'F': 4}
    return [DAY_MAP[c] for c in s if c in DAY_MAP]


def _parse_time(t):
    """'8:30am' → 510"""
    m = re.match(r'(\d+):(\d+)\s*(am|pm)', t.strip(), re.IGNORECASE)
    if not m:
        return None
    h, mn, p = int(m.group(1)), int(m.group(2)), m.group(3).lower()
    if p == 'pm' and h != 12:
        h += 12
    elif p == 'am' and h == 12:
        h = 0
    return h * 60 + mn


def _parse_time_range(t):
    """'8:30am-9:25am' → (510, 565)"""
    parts = t.split('-')
    if len(parts) != 2:
        return None, None
    return _parse_time(parts[0].strip()), _parse_time(parts[1].strip())


def _is_regular(date_str):
    """True if date range spans multiple days (not a single exam date)."""
    parts = date_str.strip().split('-')
    return len(parts) == 2 and parts[0].strip() != parts[1].strip()


def _parse_exam_date(date_str):
    """'05/12-05/12' → '05/12' (the single date portion)."""
    parts = date_str.strip().split('-')
    return parts[0].strip() if parts else None


def _exam_date_to_iso(date_part, year):
    """'05/12' + 2026 → '2026-05-12'."""
    m = re.match(r'(\d{2})/(\d{2})', date_part)
    if not m:
        return None
    return f"{year}-{m.group(1)}-{m.group(2)}"


def _parse_regular_date_range(date_str, year):
    """'01/26-05/08' + 2026 → ('2026-01-26', '2026-05-08'), or (None, None)."""
    parts = date_str.strip().split('-')
    if len(parts) != 2:
        return None, None
    start = _exam_date_to_iso(parts[0].strip(), year)
    end   = _exam_date_to_iso(parts[1].strip(), year)
    return start, end


def _extract_year_from_term(term_str):
    """Extract the calendar year for a term's Meeting Dates from a Banner-style
    term code like '202710' (YYYY + 2-digit term suffix). Fall terms (suffix
    '10') are coded under the following calendar year's label — e.g. Fall 2026
    is '202710' — so the year is decremented by one to get the true calendar
    year those Meeting Dates fall in."""
    m = re.match(r'(\d{4})(\d{2})?', str(term_str).strip())
    if not m:
        return None
    year = int(m.group(1))
    if m.group(2) == '10':
        year -= 1
    return year


# ── routes ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return send_file(os.path.join(_BASE_DIR, "static", "index.html"))


@app.route("/api/data", methods=["GET"])
def get_data():
    return jsonify(load_data())


@app.route("/api/data", methods=["POST"])
def post_data():
    if not get_data_file():
        return jsonify({"error": "No file is open — use Save As to choose one."}), 400
    data = request.get_json()
    save_data_to_file(data)
    return jsonify({"status": "ok"})


@app.route("/api/file-path", methods=["GET"])
def file_path():
    return jsonify({"path": get_data_file()})


def _default_dir():
    """Best guess for the initial directory in file dialogs."""
    current = get_data_file()
    if current:
        return os.path.dirname(current)
    return os.path.expanduser("~")


@app.route("/api/saveas", methods=["POST"])
def save_as():
    from webview import FileDialog
    data = request.get_json()
    current = get_data_file()
    default_name = os.path.basename(current) if current else "schedule.json"
    path = _file_dialog(FileDialog.SAVE,
                        save_filename=default_name,
                        file_types=("JSON files (*.json)", "All files (*.*)"))
    if path is None:
        return jsonify({"cancelled": True})
    if not path.endswith(".json"):
        path += ".json"
    set_data_file(path)
    save_data_to_file(data)
    return jsonify({"path": path})


@app.route("/api/open-dialog", methods=["POST"])
def open_dialog():
    from webview import FileDialog
    path = _file_dialog(FileDialog.OPEN,
                        file_types=("JSON files (*.json)", "All files (*.*)"))
    if path is None:
        return jsonify({"cancelled": True})
    try:
        with open(path) as f:
            data = json.load(f)
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    set_data_file(path)
    return jsonify({"path": path, "data": data})


@app.route("/api/import-csv", methods=["POST"])
def import_csv_route():
    from webview import FileDialog
    path = _file_dialog(FileDialog.OPEN,
                        file_types=("CSV files (*.csv)", "All files (*.*)"))
    if path is None:
        return jsonify({"cancelled": True})

    DAY_NAMES = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri']
    grad_courses = []
    undergrad = {}  # key: "SUBJ NUMBER" → {subject, number, title, sections:[]}
    exam_courses = {}  # key: course_name → set of (date_iso, start_min, end_min)
    term_year = None

    try:
        with open(path, newline='', encoding='utf-8-sig') as f:
            reader = csv.DictReader(f)
            for row in reader:
                level = row.get("Course Level", "").strip()
                subject = row.get("Subject", "").strip()
                number = row.get("Number", "").strip()
                section = row.get("Section", "").strip()
                title = row.get("Title", "").strip()
                days_raw = row.get("Meeting Days", "").strip()
                times_raw = row.get("Meeting Times", "").strip()
                dates_raw = row.get("Meeting Dates", "").strip()

                # Try to extract year from Term column
                if term_year is None:
                    term_val = row.get("Term", "").strip()
                    if term_val:
                        term_year = _extract_year_from_term(term_val)

                if not days_raw or not times_raw:
                    continue

                slots_days = days_raw.split("|")
                slots_times = times_raw.split("|")
                slots_dates = dates_raw.split("|") if dates_raw else []

                regular, exams = [], []
                date_start, date_end = None, None
                for i, d in enumerate(slots_days):
                    t_str = slots_times[i].strip() if i < len(slots_times) else ""
                    dt_str = slots_dates[i].strip() if i < len(slots_dates) else ""
                    days = _parse_days(d.strip())
                    s_min, e_min = _parse_time_range(t_str)
                    if not days or s_min is None:
                        continue
                    if _is_regular(dt_str):
                        for day in days:
                            regular.append({"day": day, "start_min": s_min, "end_min": e_min})
                        if date_start is None and term_year:
                            date_start, date_end = _parse_regular_date_range(dt_str, term_year)
                    else:
                        # Exam: capture the actual date
                        date_part = _parse_exam_date(dt_str)
                        date_iso = _exam_date_to_iso(date_part, term_year) if date_part and term_year else None
                        for day in days:
                            exam_entry = {"day": day, "start_min": s_min, "end_min": e_min}
                            if date_iso:
                                exam_entry["date"] = date_iso
                            elif date_part:
                                exam_entry["date_raw"] = date_part
                            exams.append(exam_entry)

                course_name = f"{subject} {number}"
                section_label = section.strip()

                # Collect exam info for exam_courses (dedup by date+time)
                if level == "Undergraduate":
                    for ex in exams:
                        iso = ex.get("date") or (
                            _exam_date_to_iso(ex["date_raw"], term_year)
                            if ex.get("date_raw") and term_year else None)
                        if iso:
                            exam_courses.setdefault(course_name, set()).add(
                                (iso, ex["start_min"], ex["end_min"]))

                if level == "Graduate":
                    if regular:
                        grad_courses.append({
                            "name": course_name,
                            "section": section_label,
                            "day": regular[0]["day"],
                            "start_min": regular[0]["start_min"],
                            "end_min": regular[0]["end_min"],
                            "meetings": regular,
                            "exams": exams,
                            "date_start": date_start,
                            "date_end":   date_end,
                        })
                elif level == "Undergraduate":
                    key = course_name
                    if key not in undergrad:
                        undergrad[key] = {"subject": subject, "number": number,
                                          "title": title, "sections": []}
                    if regular:
                        undergrad[key]["sections"].append({
                            "name": course_name,
                            "section": section_label,
                            "day": regular[0]["day"],
                            "start_min": regular[0]["start_min"],
                            "end_min": regular[0]["end_min"],
                            "meetings": regular,
                            "exams": exams,
                            "date_start": date_start,
                            "date_end":   date_end,
                        })

        # Build exam_courses response: list of {name, exams: [{date, start_min, end_min}]}
        exam_courses_list = []
        for cname in sorted(exam_courses.keys()):
            unique_exams = sorted(exam_courses[cname])
            exam_courses_list.append({
                "name": cname,
                "exams": [{"date": e[0], "start_min": e[1], "end_min": e[2]}
                          for e in unique_exams],
            })

        return jsonify({
            "grad_courses": grad_courses,
            "undergrad_courses": sorted(undergrad.values(),
                                        key=lambda c: (c["subject"], c["number"])),
            "exam_courses": exam_courses_list,
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@app.route("/api/schedule", methods=["POST"])
def run_schedule():
    data = request.get_json()
    try:
        result = solve(data)
        return jsonify(result)
    except Exception:
        traceback.print_exc()
        return jsonify({"status": "error", "message": traceback.format_exc()}), 500


@app.route("/api/schedule-proctoring", methods=["POST"])
def run_proctoring():
    data = request.get_json()
    try:
        result = solve_proctoring(data)
        return jsonify(result)
    except Exception:
        traceback.print_exc()
        return jsonify({"status": "error", "message": traceback.format_exc()}), 500


@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    from webview import FileDialog
    data = request.get_json()
    path = _file_dialog(FileDialog.SAVE,
                        save_filename="ta_schedule.docx",
                        file_types=("Word documents (*.docx)", "All files (*.*)"))
    if not path or path in ("/", ""):
        return jsonify({"cancelled": True})
    if not path.endswith(".docx"):
        path += ".docx"
    try:
        doc = generate_docx(data)
        doc.save(path)
        return jsonify({"path": path})
    except Exception:
        traceback.print_exc()
        return jsonify({"error": traceback.format_exc()}), 500


# ── solver ───────────────────────────────────────────────────────────────────

_SOLVER_ITERATIONS = 50

_BASE_SCORE          = 1000
_EXPERIENCE_BONUS    = 200
_NEW_ROLE_PENALTY    = 800   # cost of opening a new (TA, role) pairing
_LOAD_BALANCE_WEIGHT = 500

_LAB_FAMILIARITY_BONUS      = 300
_LAB_SECTION_BONUS          = 150
_SAME_COURSE_PROCTOR_BONUS  = 200
_SAME_SECTION_PROCTOR_BONUS = 100
_PROC_LOAD_BALANCE_WEIGHT   = 500

_SPREAD_WINDOW_DAYS     = 7    # exams within this many days of each other count as "clustered"
_SPREAD_PENALTY_PER_DAY = 40   # max penalty 280 (same-day) — under one PE load-balance unit (500),
                                # so this nudges ties rather than overriding load balance or
                                # course/section familiarity

# Single source of truth for user-configurable solver weights, persisted in
# data["settings"]. Keep keys in sync with DEFAULT_SETTINGS in static/index.html.
_DEFAULT_SETTINGS = {
    "experience_bonus": _EXPERIENCE_BONUS,
    "new_role_penalty": _NEW_ROLE_PENALTY,
    "load_balance_weight": _LOAD_BALANCE_WEIGHT,
    "lab_familiarity_bonus": _LAB_FAMILIARITY_BONUS,
    "lab_section_bonus": _LAB_SECTION_BONUS,
    "same_course_proctor_bonus": _SAME_COURSE_PROCTOR_BONUS,
    "same_section_proctor_bonus": _SAME_SECTION_PROCTOR_BONUS,
    "proc_load_balance_weight": _PROC_LOAD_BALANCE_WEIGHT,
    "spread_window_days": _SPREAD_WINDOW_DAYS,
    "spread_penalty_per_day": _SPREAD_PENALTY_PER_DAY,
}


def _ta_date_conflict_blocks_lab(ta, lab, meetings):
    """True if any of the TA's date_conflicts[] (not marked ignore_for_labs)
    overlaps this lab's meetings once clamped to the lab's own date_start/
    date_end (or the conflict's own span, if the lab has no date range)."""
    ds, de = lab.get("date_start"), lab.get("date_end")
    for dc in ta.get("date_conflicts", []):
        if dc.get("ignore_for_labs"):
            continue
        for wd, sm, em in _expand_date_conflict_weekdays(dc, ds, de):
            if any(ld == wd and times_overlap(sm, em, ls, le) for ld, ls, le in meetings):
                return True
    return False


def solve(data):
    roles_map = {r["id"]: r for r in data.get("roles", [])}
    labs      = data.get("labs", [])
    tas       = data.get("tas", [])
    assignments_in = data.get("assignments", [])
    gc_map    = {gc["id"]: gc for gc in data.get("grad_courses", [])}

    settings = {**_DEFAULT_SETTINGS, **(data.get("settings") or {})}
    experience_bonus    = settings["experience_bonus"]
    new_role_penalty    = settings["new_role_penalty"]
    load_balance_weight = settings["load_balance_weight"]

    if not labs or not tas:
        return {"status": "feasible", "assignments": assignments_in, "diagnostics": {}}

    labs_by_id = {lab["id"]: lab for lab in labs}
    locked_assignments = [a for a in assignments_in if a.get("locked")]

    # ── precomputed, input-only data (constant across all iterations) ──────

    lab_meetings = {lab["id"]: _get_meetings(lab) for lab in labs}

    # Every fixed time block a TA already owns: grad courses + other commitments.
    ta_fixed_times = {}
    for ta in tas:
        fixed = []
        for gc_id in ta.get("grad_course_ids", []):
            gc = gc_map.get(gc_id)
            if gc:
                fixed.extend(_get_meetings(gc))
        for oc in ta.get("other_commitments", []):
            fixed.append((oc["day"], oc["start_min"], oc["end_min"]))
        ta_fixed_times[ta["id"]] = fixed

    # (ta_id, lab_id) pairs blocked by a fixed conflict. Depends only on the
    # input, so it is built once rather than re-derived on every eligibility test.
    static_conflicts = frozenset(
        (ta["id"], lab["id"])
        for ta in tas
        for lab in labs
        if any(ld == fd and times_overlap(ls, le, fs, fe)
               for ld, ls, le in lab_meetings[lab["id"]]
               for fd, fs, fe in ta_fixed_times[ta["id"]])
        or _ta_date_conflict_blocks_lab(ta, lab, lab_meetings[lab["id"]])
    )

    # ── build work list (same every iteration) ─────────────────────────────

    # Initialise locked state to compute the slot list and initial eligibility
    init_ta_per_slot = {}
    for a in locked_assignments:
        key = (a["lab_id"], a["role_id"])
        init_ta_per_slot.setdefault(key, set()).add(a["ta_id"])

    slots = []
    for lab in labs:
        for rr in lab.get("roles", []):
            count        = rr.get("count", 1)
            locked_count = len(init_ta_per_slot.get((lab["id"], rr["role_id"]), set()))
            for _ in range(count - locked_count):
                slots.append((lab, rr))

    if not slots:
        # Nothing to assign — return locked assignments as-is
        return {"status": "feasible", "assignments": locked_assignments, "diagnostics": {}}

    # ── mutable per-pass state ─────────────────────────────────────────────

    def initial_state():
        """Fresh solver state seeded from the locked assignments."""
        st = {"used_se": {}, "booked_labs": {}, "roles": {}, "per_slot": {}}
        for ta in tas:
            tid = ta["id"]
            st["used_se"][tid]     = sum(d.get("se_value", 0)
                                         for d in ta.get("outside_duties", []))
            st["booked_labs"][tid] = set()
            st["roles"][tid]       = set()
        for a in locked_assignments:
            tid, lid, rid = a["ta_id"], a["lab_id"], a["role_id"]
            lab  = labs_by_id.get(lid)
            role = roles_map.get(rid, {})
            if lab and tid in st["used_se"]:
                st["used_se"][tid] += role.get("se_value", 1.0)
                st["booked_labs"][tid].add(lid)
                st["roles"][tid].add(rid)
            st["per_slot"].setdefault((lid, rid), set()).add(tid)
        return st

    def double_booked(st, ta_id, lab_mtgs):
        for booked_id in st["booked_labs"][ta_id]:
            for bd, bs, be in lab_meetings.get(booked_id, ()):
                for ld, ls, le in lab_mtgs:
                    if ld == bd and times_overlap(ls, le, bs, be):
                        return True
        return False

    def eligible_tas(st, lab, rr):
        role     = roles_map.get(rr["role_id"], {})
        se_val   = role.get("se_value", 1.0)
        lab_id   = lab["id"]
        lab_mtgs = lab_meetings[lab_id]
        already  = st["per_slot"].get((lab_id, rr["role_id"]), set())
        result   = []
        for ta in tas:
            tid = ta["id"]
            if tid in already:
                continue
            if (tid, lab_id) in static_conflicts:
                continue
            if st["used_se"][tid] + se_val > ta.get("max_se", 2.0) + 0.001:
                continue
            if double_booked(st, tid, lab_mtgs):
                continue
            result.append(ta)
        return result

    def score(st, ta, lab, rr):
        s = _BASE_SCORE
        if rr.get("preferred_experienced", 0) > 0 and ta.get("experience") == "experienced":
            s += experience_bonus
        roles_held = st["roles"][ta["id"]]
        if rr["role_id"] not in roles_held:
            # Opening a new (TA, role) pairing is the thing being minimised, so it is
            # charged even for a TA's first role — that is what lets a TA already in
            # this role outrank an idle one. Scaling by roles already held sends
            # unavoidable new pairings to the least-fragmented TA.
            s -= new_role_penalty * (1 + len(roles_held))
        s -= st["used_se"][ta["id"]] * load_balance_weight
        s += random.random()
        return s

    # Slot order is derived from locked-only state, so it is identical on every
    # iteration — compute it once.
    init_st = initial_state()
    sorted_slots = sorted(slots,
        key=lambda s: (len(eligible_tas(init_st, s[0], s[1])),
                       -roles_map.get(s[1]["role_id"], {}).get("se_value", 1.0)))

    # ── single greedy pass ─────────────────────────────────────────────────

    def _greedy_pass():
        st = initial_state()
        result_assignments = list(locked_assignments)

        for lab, rr in sorted_slots:
            candidates = eligible_tas(st, lab, rr)
            if not candidates:
                continue
            best = max(candidates, key=lambda ta: score(st, ta, lab, rr))
            role = roles_map.get(rr["role_id"], {})

            result_assignments.append({
                "lab_id":  lab["id"],
                "role_id": rr["role_id"],
                "ta_id":   best["id"],
                "locked":  False,
            })

            st["used_se"][best["id"]] += role.get("se_value", 1.0)
            st["booked_labs"][best["id"]].add(lab["id"])
            st["roles"][best["id"]].add(rr["role_id"])
            st["per_slot"].setdefault((lab["id"], rr["role_id"]), set()).add(best["id"])

        return result_assignments

    # ── run multiple iterations, keep the best ─────────────────────────────

    best_result = None
    best_unfilled = float("inf")

    for _ in range(_SOLVER_ITERATIONS):
        result_assignments = _greedy_pass()

        # Count unfilled seats for this attempt
        unfilled_count = len(slots) - sum(
            1 for a in result_assignments if not a.get("locked"))
        if unfilled_count <= 0:
            best_result = result_assignments
            break
        if unfilled_count < best_unfilled:
            best_unfilled = unfilled_count
            best_result = result_assignments

    result_assignments = best_result

    # ── diagnostics ──────────────────────────────────────────────────────────
    tas_map = {ta["id"]: ta for ta in tas}
    unfilled, unfulfilled_exp = [], []
    for lab in labs:
        for rr in lab.get("roles", []):
            role_id    = rr["role_id"]
            count      = rr.get("count", 1)
            pref_exp   = rr.get("preferred_experienced", 0)
            role_label = roles_map.get(role_id, {}).get("label", role_id)
            role_asgn  = [a for a in result_assignments
                          if a["lab_id"] == lab["id"] and a["role_id"] == role_id]
            if len(role_asgn) < count:
                unfilled.append({
                    "lab_name":   lab["name"],
                    "role_label": role_label,
                    "assigned":   len(role_asgn),
                    "required":   count,
                })
            if pref_exp > 0:
                exp_n = sum(1 for a in role_asgn
                            if tas_map.get(a["ta_id"], {}).get("experience") == "experienced")
                if exp_n < pref_exp:
                    unfulfilled_exp.append({
                        "lab_name":      lab["name"],
                        "role_label":    role_label,
                        "exp_assigned":  exp_n,
                        "exp_preferred": pref_exp,
                    })

    status = "partial" if (unfilled or unfulfilled_exp) else "feasible"
    return {
        "status":      status,
        "assignments": result_assignments,
        "diagnostics": {"unfilled_roles": unfilled, "unfulfilled_experience": unfulfilled_exp},
    }


# ── proctoring solver ────────────────────────────────────────────────────────

def _no_fixed_time(exam):
    """True if an exam's own start_min/end_min are unknown — either the whole
    date is TBD, or the date is known but the time isn't."""
    return bool(exam.get("time_tbd") or exam.get("tbd"))


def solve_proctoring(data):
    exams = data.get("exams", [])
    tas = data.get("tas", [])
    proctor_in = data.get("proctor_assignments", [])
    labs = data.get("labs", [])
    gc_map = {gc["id"]: gc for gc in data.get("grad_courses", [])}
    assignments = data.get("assignments", [])

    settings = {**_DEFAULT_SETTINGS, **(data.get("settings") or {})}
    lab_familiarity_bonus      = settings["lab_familiarity_bonus"]
    lab_section_bonus          = settings["lab_section_bonus"]
    same_course_proctor_bonus  = settings["same_course_proctor_bonus"]
    same_section_proctor_bonus = settings["same_section_proctor_bonus"]
    proc_load_balance_weight   = settings["proc_load_balance_weight"]
    spread_window_days         = settings["spread_window_days"]
    spread_penalty_per_day     = settings["spread_penalty_per_day"]

    if not exams or not tas:
        return {"status": "feasible", "proctor_assignments": proctor_in, "diagnostics": {}}

    exams_by_id = {ex["id"]: ex for ex in exams}
    locked = [a for a in proctor_in if a.get("locked")]

    # Build lab assignments per TA: ta_id → [(weekday, start_min, end_min)]
    labs_by_id = {l["id"]: l for l in labs}
    ta_lab_times = {}
    for a in assignments:
        lab = labs_by_id.get(a["lab_id"])
        if not lab:
            continue
        for d, s, e in _get_meetings(lab):
            ta_lab_times.setdefault(a["ta_id"], []).append((d, s, e))

    # TA → set of lab course names and sections (for familiarity bonus)
    ta_lab_courses = {}
    ta_lab_sections = {}
    for a in assignments:
        lab = labs_by_id.get(a["lab_id"])
        if lab:
            ta_lab_courses.setdefault(a["ta_id"], set()).add(lab.get("name", ""))
            sect = lab.get("section", "")
            if sect:
                ta_lab_sections.setdefault(a["ta_id"], set()).add((lab.get("name", ""), sect))

    def _parse_exam_date(exam):
        try:
            return datetime.date.fromisoformat(exam["date"])
        except (KeyError, ValueError):
            return None

    # Build slots
    init_per_exam = {}
    for a in locked:
        init_per_exam.setdefault(a["exam_id"], set()).add(a["ta_id"])

    slots = []
    for exam in exams:
        count = exam.get("proctor_count", 1)
        locked_count = len(init_per_exam.get(exam["id"], set()))
        for _ in range(count - locked_count):
            slots.append(exam)

    if not slots:
        return {"status": "feasible", "proctor_assignments": locked, "diagnostics": {}}

    # ── precomputed, input-only data (constant across all iterations) ──────

    slot_exams = {ex["id"]: ex for ex in slots}
    # Covers every exam (not just slot_exams) since initial_state() seeds
    # st["times"] from locked assignments, which may reference exams that are
    # fully locked and therefore absent from slots/slot_exams.
    exam_date_obj = {eid: _parse_exam_date(ex) for eid, ex in exams_by_id.items()}
    exam_wd = {eid: (exam_date_obj[eid].weekday() if exam_date_obj[eid] else None)
               for eid in slot_exams}

    def _fixed_conflict(ta, exam):
        """True if a TA has an input-derived conflict with this exam: an assigned
        lab, a grad course meeting or exam, a commitment, or a date conflict."""
        eid = exam["id"]
        exam_dobj = exam_date_obj.get(eid)

        if exam.get("time_tbd"):
            # The exam's own time is unknown, so weekday/time-based checks (lab,
            # grad course, commitments) and partial-day date_conflicts can't be
            # evaluated. A date_conflicts[] entry that blocks the TA's ENTIRE
            # day still applies, though — no matter what time the exam ends up
            # at, it would fall inside it.
            if exam_dobj is not None:
                for dc in ta.get("date_conflicts", []):
                    days = _date_conflict_days(dc)
                    if days is None:
                        continue
                    if _date_conflict_window_for_day(dc, days, exam_dobj) == (0, 1440):
                        return True
            return False

        if exam.get("tbd"):
            return False  # no date at all — nothing left to check

        wd = exam_wd.get(exam["id"])
        es, ee = exam.get("start_min", 0), exam.get("end_min", 0)
        if wd is not None:
            for ld, ls, le in ta_lab_times.get(ta["id"], []):
                if ld == wd and times_overlap(es, ee, ls, le):
                    return True
            for gc_id in ta.get("grad_course_ids", []):
                gc = gc_map.get(gc_id)
                if not gc:
                    continue
                for gd, gs, ge in _get_meetings(gc):
                    if gd == wd and times_overlap(es, ee, gs, ge):
                        return True
                exam_date = exam.get("date")
                if exam_date:
                    for gc_ex in gc.get("exams", []):
                        if gc_ex.get("date") == exam_date and times_overlap(
                                es, ee, gc_ex.get("start_min", 0), gc_ex.get("end_min", 0)):
                            return True
            for oc in ta.get("other_commitments", []):
                if oc["day"] == wd and times_overlap(
                        es, ee, oc["start_min"], oc["end_min"]):
                    return True
        if exam_dobj is not None:
            for dc in ta.get("date_conflicts", []):
                if _date_conflict_overlaps(dc, exam_dobj, es, ee):
                    return True
        return False

    static_conflicts = frozenset(
        (ta["id"], eid)
        for ta in tas
        for eid, exam in slot_exams.items()
        if _fixed_conflict(ta, exam)
    )

    # ── mutable per-pass state ─────────────────────────────────────────────

    def initial_state():
        """Fresh solver state seeded from the locked proctor assignments."""
        st = {"used_pe": {}, "assigned_exams": {}, "times": {},
              "courses": {}, "sections": {}}
        for ta in tas:
            tid = ta["id"]
            st["used_pe"][tid]        = sum(op.get("pe_value", 0)
                                            for op in ta.get("outside_proctoring", []))
            st["assigned_exams"][tid] = set()
            st["times"][tid]          = []
            st["courses"][tid]        = set()
            st["sections"][tid]       = set()
        for a in locked:
            tid, eid = a["ta_id"], a["exam_id"]
            exam = exams_by_id.get(eid)
            if not exam or tid not in st["used_pe"]:
                continue
            st["used_pe"][tid] += exam.get("pe_value", 1.0)
            st["assigned_exams"][tid].add(eid)
            if not _no_fixed_time(exam):
                st["times"][tid].append(
                    (exam_date_obj.get(eid), exam.get("start_min", 0), exam.get("end_min", 0)))
            course_name = exam.get("course_name", "")
            sect = exam.get("section", "")
            if course_name:
                st["courses"][tid].add(course_name)
            if course_name and sect:
                st["sections"][tid].add((course_name, sect))
        return st

    def eligible_tas(st, exam):
        pe_val = exam.get("pe_value", 1.0)
        eid    = exam["id"]
        edate  = exam_date_obj.get(eid)
        es, ee = exam.get("start_min", 0), exam.get("end_min", 0)
        result = []
        for ta in tas:
            tid = ta["id"]
            if st["used_pe"][tid] + pe_val > ta.get("max_pe", 2.0) + 0.001:
                continue
            if eid in st["assigned_exams"][tid]:
                continue
            # Same-date time conflicts with exams already being proctored
            if not _no_fixed_time(exam) and any(
                    pdate == edate and times_overlap(es, ee, ps, pe)
                    for pdate, ps, pe in st["times"][tid]):
                continue
            if (tid, eid) in static_conflicts:
                continue
            result.append(ta)
        return result

    def score(st, ta, exam):
        s = 1000
        tid = ta["id"]
        course_name = exam.get("course_name", "")
        section = exam.get("section", "")
        key = (course_name, section)
        # Lab familiarity bonus (same course)
        if course_name and course_name in ta_lab_courses.get(tid, set()):
            s += lab_familiarity_bonus
        # Lab section bonus (same course + section)
        if course_name and section and key in ta_lab_sections.get(tid, set()):
            s += lab_section_bonus
        # Same-course proctoring bonus
        if course_name and course_name in st["courses"][tid]:
            s += same_course_proctor_bonus
        # Same-section proctoring bonus
        if course_name and section and key in st["sections"][tid]:
            s += same_section_proctor_bonus
        # Spread bonus — discourage clustering this TA's exams close together in time
        if not _no_fixed_time(exam):
            exam_date = exam_date_obj.get(exam["id"])
            if exam_date is not None:
                prior_dates = [d for d, _, _ in st["times"][tid] if d is not None]
                if prior_dates:
                    min_gap = min(abs((exam_date - d).days) for d in prior_dates)
                    if min_gap < spread_window_days:
                        s -= (spread_window_days - min_gap) * spread_penalty_per_day
        # Load balancing
        s -= st["used_pe"][tid] * proc_load_balance_weight
        s += random.random()
        return s

    # Per-exam eligible count is derived from locked-only state, so it is
    # identical on every iteration — compute it once. The tie-break order among
    # slots sharing a count is randomized per iteration (see _greedy_pass), so
    # tbd/time_tbd exams — which tend to tie for the highest eligible count and
    # would otherwise always sort last, since new exams are appended to the end
    # of data["exams"] — aren't deterministically starved on every pass.
    init_st = initial_state()
    eligible_count = {ex["id"]: len(eligible_tas(init_st, ex)) for ex in slot_exams.values()}

    def _greedy_pass():
        st = initial_state()
        order = sorted(slots, key=lambda ex: (eligible_count[ex["id"]], random.random()))
        result = list(locked)
        for exam in order:
            candidates = eligible_tas(st, exam)
            if not candidates:
                continue
            best = max(candidates, key=lambda ta: score(st, ta, exam))
            tid = best["id"]
            result.append({
                "exam_id": exam["id"],
                "ta_id": tid,
                "locked": False,
            })
            st["used_pe"][tid] += exam.get("pe_value", 1.0)
            st["assigned_exams"][tid].add(exam["id"])
            if not _no_fixed_time(exam):
                st["times"][tid].append(
                    (exam_date_obj.get(exam["id"]), exam.get("start_min", 0), exam.get("end_min", 0)))
            cname = exam.get("course_name", "")
            sect = exam.get("section", "")
            if cname:
                st["courses"][tid].add(cname)
            if cname and sect:
                st["sections"][tid].add((cname, sect))

        return result

    best_result = None
    best_unfilled = float("inf")

    for _ in range(_SOLVER_ITERATIONS):
        result = _greedy_pass()
        unfilled_count = len(slots) - sum(1 for a in result if not a.get("locked"))
        if unfilled_count <= 0:
            best_result = result
            break
        if unfilled_count < best_unfilled:
            best_unfilled = unfilled_count
            best_result = result

    result = best_result

    # Diagnostics
    unfilled = []
    for exam in exams:
        count = exam.get("proctor_count", 1)
        assigned = [a for a in result if a["exam_id"] == exam["id"]]
        if len(assigned) < count:
            unfilled.append({
                "exam_name": exam.get("name", ""),
                "assigned": len(assigned),
                "required": count,
            })

    status = "partial" if unfilled else "feasible"
    return {
        "status": status,
        "proctor_assignments": result,
        "diagnostics": {"unfilled_proctors": unfilled},
    }


# ── DOCX export ──────────────────────────────────────────────────────────────

def _style_docx(doc):
    """Swap the default Cambria/Calibri theme for a cleaner, more modern pair
    and give the heading levels enough size/weight/color contrast to read as
    a hierarchy on their own, since the document is mostly headings and
    tables with very little body text."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    HEADING_FONT = "Trebuchet MS"
    BODY_FONT = "Arial"
    INK = RGBColor(0x1F, 0x29, 0x37)
    ACCENT = RGBColor(0x2F, 0x54, 0x96)
    MUTED = RGBColor(0x59, 0x59, 0x59)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    specs = {
        # style name: (size, bold, italic, color, all_caps, space_before, space_after)
        "Title":     (26, True,  False, INK,    False, 0,  14),
        "Heading 1": (15, True,  False, ACCENT, True,  20, 4),
        "Heading 2": (14, True,  False, INK,    False, 10, 4),
        "Heading 3": (11.5, True, False, MUTED, False, 8,  2),
        "Heading 4": (10.5, False, True, MUTED, False, 6,  2),
    }
    for name, (size, bold, italic, color, all_caps, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = HEADING_FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = color
        style.font.all_caps = all_caps
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    # A hairline rule under Heading 1 marks the document's three major parts.
    pPr = doc.styles["Heading 1"].element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2F5496")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx(data):
    from docx import Document

    doc = Document()
    _style_docx(doc)
    doc.add_heading("TA Schedule", 0)

    roles_map = {r["id"]: r for r in data.get("roles", [])}
    tas_map = {t["id"]: t for t in data.get("tas", [])}
    labs_map = {l["id"]: l for l in data.get("labs", [])}
    gc_map = {gc["id"]: gc for gc in data.get("grad_courses", [])}
    assignments = data.get("assignments", [])
    exams_map = {e["id"]: e for e in data.get("exams", [])}
    proctor_assignments = data.get("proctor_assignments", [])

    def lab_disp(lab):
        s = lab.get("section", "")
        return f"{lab['name']} {s}".strip() if s else lab["name"]

    def full_exam_label(exam):
        cname = exam.get("course_name", "")
        sect = exam.get("section", "")
        label = f"{cname} {sect}".strip() if cname else exam.get("name", "Exam")
        sub = exam.get("name", "")
        return f"{label} — {sub}" if sub and label != sub else label or sub

    def fill_cell_lines(cell, lines):
        cell.paragraphs[0].text = lines[0] if lines else ""
        for line in lines[1:]:
            cell.add_paragraph(line)

    # Lab-centric
    doc.add_heading("Lab Assignments", 1)
    labs = data.get("labs", [])
    course_names = sorted({lab.get("name", "") for lab in labs})
    for course_name in course_names:
        course_labs = sorted(
            [lab for lab in labs if lab.get("name", "") == course_name],
            key=lambda l: l.get("section", ""),
        )
        doc.add_heading(course_name, 2)
        course_lab_ids = {lab["id"] for lab in course_labs}
        course_asgn = [a for a in assignments if a["lab_id"] in course_lab_ids]
        if course_asgn:
            by_ta = {}
            for a in course_asgn:
                by_ta.setdefault(a["ta_id"], []).append(a)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "TA", "Sections", "Email"
            for ta_id in sorted(by_ta, key=lambda tid: tas_map.get(tid, {}).get("name", "")):
                ta_asgns = sorted(
                    by_ta[ta_id],
                    key=lambda a: labs_map.get(a["lab_id"], {}).get("section", ""),
                )
                row = tbl.add_row().cells
                row[0].text = tas_map.get(ta_id, {}).get("name", ta_id)
                fill_cell_lines(row[1], [
                    f"{labs_map.get(a['lab_id'], {}).get('section', '')} "
                    f"({roles_map.get(a['role_id'], {}).get('label', a.get('role_id', ''))})"
                    for a in ta_asgns
                ])
                row[2].text = tas_map.get(ta_id, {}).get("email", "")
        doc.add_paragraph()
        for lab in course_labs:
            doc.add_heading(lab_disp(lab), 3)
            doc.add_paragraph(fmt_meetings(lab, DAY_LONG))
            lab_asgn = sorted(
                [a for a in assignments if a["lab_id"] == lab["id"]],
                key=lambda a: tas_map.get(a["ta_id"], {}).get("name", ""),
            )
            if lab_asgn:
                tbl = doc.add_table(rows=1, cols=3)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text = "Role", "TA", "Email"
                for a in lab_asgn:
                    row = tbl.add_row().cells
                    row[0].text = roles_map.get(a["role_id"], {}).get("label", a.get("role_id", ""))
                    row[1].text = tas_map.get(a["ta_id"], {}).get("name", a.get("ta_id", ""))
                    row[2].text = tas_map.get(a["ta_id"], {}).get("email", "")
            else:
                doc.add_paragraph("No assignments")
            doc.add_paragraph()

    other_duty_rows = [
        (ta, od)
        for ta in data.get("tas", [])
        for od in ta.get("outside_duties", [])
    ]
    if other_duty_rows:
        doc.add_heading("Other Duties", 2)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text = "TA", "Duty"
        for ta, od in sorted(other_duty_rows, key=lambda pair: (pair[0].get("name", ""), pair[1].get("label", ""))):
            row = tbl.add_row().cells
            row[0].text = ta.get("name", "")
            row[1].text = od.get("label", "Other Duty")
        doc.add_paragraph()

    # Exam proctoring
    exams = data.get("exams", [])
    if exams and proctor_assignments:
        doc.add_heading("Exam Proctoring", 1)

        def exam_course_key(exam):
            return exam.get("course_name") or exam.get("name", "")

        course_keys = sorted({exam_course_key(e) for e in exams})
        for course_key in course_keys:
            course_exams = [e for e in exams if exam_course_key(e) == course_key]
            course_exam_ids = {e["id"] for e in course_exams}
            if not any(a["exam_id"] in course_exam_ids for a in proctor_assignments):
                continue
            doc.add_heading(course_key or "Exam", 2)
            sections = sorted({e.get("section", "") for e in course_exams})
            for section in sections:
                section_exams = sorted(
                    [e for e in course_exams if e.get("section", "") == section],
                    key=lambda e: e.get("date", ""),
                )
                section_exam_ids = {e["id"] for e in section_exams}
                if not any(a["exam_id"] in section_exam_ids for a in proctor_assignments):
                    continue
                section_heading = f"{course_key} {section}".strip() if course_key else section
                doc.add_heading(section_heading or "—", 3)
                for exam in section_exams:
                    exam_asgn = sorted(
                        [a for a in proctor_assignments if a["exam_id"] == exam["id"]],
                        key=lambda a: tas_map.get(a["ta_id"], {}).get("name", ""),
                    )
                    if not exam_asgn:
                        continue
                    if exam.get("tbd"):
                        heading = f"{exam.get('name', 'Exam')} — Date/Time TBD"
                    else:
                        if exam.get("time_tbd"):
                            time_str = "Time TBD"
                        else:
                            time_str = (
                                f"{fmt_time(exam.get('start_min', 0))} – "
                                f"{fmt_time(exam.get('end_min', 0))}"
                            )
                        heading = f"{exam.get('name', 'Exam')} — {exam.get('date', '—')}, {time_str}"
                    doc.add_heading(heading, 4)
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.style = "Table Grid"
                    hdr = tbl.rows[0].cells
                    hdr[0].text, hdr[1].text = "TA", "Email"
                    for a in exam_asgn:
                        ta = tas_map.get(a["ta_id"])
                        row = tbl.add_row().cells
                        row[0].text = ta.get("name", a["ta_id"]) if ta else a["ta_id"]
                        row[1].text = ta.get("email", "") if ta else ""
            doc.add_paragraph()

    # TA-centric
    doc.add_heading("TA Assignments", 1)
    for ta in sorted(data.get("tas", []), key=lambda t: t.get("name", "")):
        doc.add_heading(ta["name"], 2)
        if ta.get("email"):
            doc.add_paragraph(f"Email: {ta['email']}")
        schedule_lines = []
        for gc_id in ta.get("grad_course_ids", []):
            gc = gc_map.get(gc_id)
            if gc:
                schedule_lines.append(f"{lab_disp(gc)} — {fmt_meetings(gc, DAY_SHORT)}")
        for oc in ta.get("other_commitments", []):
            schedule_lines.append(
                f"{oc['label']} — {DAY_SHORT[oc['day']]} {fmt_time(oc['start_min'])}–{fmt_time(oc['end_min'])}"
            )
        for dc in ta.get("date_conflicts", []):
            start_date, end_date = _dc_dates(dc)
            start_min, end_min = dc.get("start_min", 0), dc.get("end_min", 0)
            if start_date == end_date:
                date_str = f"{start_date} {fmt_time(start_min)}–{fmt_time(end_min)}"
            else:
                date_str = f"{start_date} {fmt_time(start_min)} – {end_date} {fmt_time(end_min)}"
            schedule_lines.append(f"{dc.get('label', '')} — {date_str}")
        if schedule_lines:
            doc.add_heading("Schedule", 3)
            for line in schedule_lines:
                doc.add_paragraph(line, style="List Bullet")
        ta_asgn = [a for a in assignments if a["ta_id"] == ta["id"]]
        outside = ta.get("outside_duties", [])
        ta_proctor = sorted(
            [a for a in proctor_assignments if a["ta_id"] == ta["id"]],
            key=lambda a: (
                exams_map.get(a["exam_id"], {}).get("date", ""),
                exams_map.get(a["exam_id"], {}).get("start_min", 0),
            ),
        )
        if ta_asgn or outside:
            doc.add_heading("Lab Assignments", 3)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Lab", "Role", "Time"
            for a in sorted(
                ta_asgn,
                key=lambda a: (
                    labs_map.get(a["lab_id"], {}).get("day", 0),
                    labs_map.get(a["lab_id"], {}).get("start_min", 0),
                ),
            ):
                lab = labs_map.get(a["lab_id"], {})
                role = roles_map.get(a["role_id"], {})
                row = tbl.add_row().cells
                row[0].text = lab_disp(lab)
                row[1].text = role.get("label", "")
                if lab:
                    row[2].text = fmt_meetings(lab)
            for od in outside:
                row = tbl.add_row().cells
                row[0].text = od.get("label", "Other Duty")
                row[1].text = "Other Duty"
                row[2].text = "—"
        outside_proctor = ta.get("outside_proctoring", [])
        if ta_proctor or outside_proctor:
            doc.add_heading("Proctoring", 3)
            ptbl = doc.add_table(rows=1, cols=3)
            ptbl.style = "Table Grid"
            phdr = ptbl.rows[0].cells
            phdr[0].text, phdr[1].text, phdr[2].text = "Date", "Time", "Exam"
            for pa in ta_proctor:
                exam = exams_map.get(pa["exam_id"], {})
                prow = ptbl.add_row().cells
                prow[0].text = exam.get("date", "—")
                prow[1].text = (
                    f"{fmt_time(exam.get('start_min', 0))}–{fmt_time(exam.get('end_min', 0))}"
                    if exam.get("date") else "—"
                )
                prow[2].text = full_exam_label(exam)
            for op in outside_proctor:
                orow = ptbl.add_row().cells
                orow[0].text = "—"
                orow[1].text = "—"
                orow[2].text = op.get("label", "Outside Proctoring")
        if not ta_asgn and not outside and not ta_proctor and not outside_proctor:
            doc.add_paragraph("No assignments")
        doc.add_paragraph()

    return doc


def _find_free_port():
    with socket.socket() as s:
        s.bind(("", 0))
        return s.getsockname()[1]


def _wait_for_flask(port, timeout=10):
    import time
    import urllib.request
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            urllib.request.urlopen(f"http://localhost:{port}/", timeout=1)
            return True
        except Exception:
            time.sleep(0.05)
    return False


if __name__ == "__main__":

    try:
        import webview
    except ImportError:
        # pywebview not installed — fall back to plain Flask + browser
        import webbrowser
        webbrowser.open("http://localhost:5050")
        app.run(debug=False, port=5050)
    else:
        port = _find_free_port()

        flask_thread = threading.Thread(
            target=lambda: app.run(host="127.0.0.1", port=port, debug=False),
            daemon=True,
        )
        flask_thread.start()
        _wait_for_flask(port)

        _window = webview.create_window(
            "TA Scheduler",
            f"http://127.0.0.1:{port}/",
            width=1500,
            height=960,
            min_size=(900, 600),
            text_select=True,
        )

        _closing_confirmed = False

        def _confirm_close():
            global _closing_confirmed
            try:
                dirty = bool(_window.evaluate_js('window.S && window.S.dirty'))
            except Exception:
                dirty = False

            if dirty and not _window.create_confirmation_dialog(
                "Unsaved Changes",
                "You have unsaved changes that have not been saved. Quit anyway?",
            ):
                return  # user chose Cancel — leave the window open

            _closing_confirmed = True
            _window.destroy()

        def _on_closing():
            if _closing_confirmed:
                return  # already confirmed (or nothing to confirm) — let this close proceed
            threading.Thread(target=_confirm_close, daemon=True).start()
            return False  # cancel *this* close attempt; _confirm_close decides the real outcome

        _window.events.closing += _on_closing

        webview.start()
