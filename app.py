import io
import json
import os
import subprocess
import traceback

from flask import Flask, jsonify, make_response, request, send_file

app = Flask(__name__, static_folder="static")

# None until the user opens or saves a file.
_data_file = None

EMPTY_DATA = {"roles": [], "grad_courses": [], "labs": [], "tas": [], "assignments": []}


def get_data_file():
    return _data_file


def set_data_file(path):
    global _data_file
    _data_file = path


# ── helpers ──────────────────────────────────────────────────────────────────

def load_data():
    p = get_data_file()
    if not p or not os.path.exists(p):
        return EMPTY_DATA
    with open(p) as f:
        return json.load(f)


def save_data_to_file(data):
    with open(get_data_file(), "w") as f:
        json.dump(data, f, indent=2)


def _osascript_dialog(script):
    """Run an AppleScript file-dialog and return the chosen path, or None."""
    try:
        r = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=120,
        )
        if r.returncode == 0:
            return r.stdout.strip() or None
    except Exception:
        pass
    return None


def times_overlap(s1, e1, s2, e2):
    return s1 < e2 and e1 > s2


def fmt_time(minutes):
    h, m = divmod(int(minutes), 60)
    ampm = "PM" if h >= 12 else "AM"
    h12 = h - 12 if h > 12 else (12 if h == 0 else h)
    return f"{h12}:{m:02d} {ampm}"


# ── routes ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return send_file("static/index.html")


@app.route("/api/data", methods=["GET"])
def get_data():
    return jsonify(load_data())


@app.route("/api/data", methods=["POST"])
def post_data():
    data = request.get_json()
    save_data_to_file(data)
    return jsonify({"status": "ok"})


@app.route("/api/file-path", methods=["GET"])
def file_path():
    return jsonify({"path": get_data_file()})


def _default_dir():
    """Best guess for the initial directory in file dialogs."""
    current = get_data_file()
    if current:
        return os.path.dirname(current)
    return os.path.expanduser("~")


@app.route("/api/saveas", methods=["POST"])
def save_as():
    data = request.get_json()
    current = get_data_file()
    default_name = os.path.basename(current) if current else "schedule.json"
    script = (
        f'set f to choose file name '
        f'with prompt "Save Schedule As" '
        f'default name "{default_name}" '
        f'default location POSIX file "{_default_dir()}"\n'
        f'return POSIX path of f'
    )
    path = _osascript_dialog(script)
    if path is None:
        return jsonify({"cancelled": True})
    if not path.endswith(".json"):
        path += ".json"
    set_data_file(path)
    save_data_to_file(data)
    return jsonify({"path": path})


@app.route("/api/open-dialog", methods=["POST"])
def open_dialog():
    script = (
        f'set f to choose file '
        f'with prompt "Open Schedule" '
        f'default location POSIX file "{_default_dir()}"\n'
        f'return POSIX path of f'
    )
    path = _osascript_dialog(script)
    if path is None:
        return jsonify({"cancelled": True})
    try:
        with open(path) as f:
            data = json.load(f)
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    set_data_file(path)
    return jsonify({"path": path, "data": data})


@app.route("/api/schedule", methods=["POST"])
def run_schedule():
    data = request.get_json()
    try:
        result = solve(data)
        return jsonify(result)
    except Exception:
        traceback.print_exc()
        return jsonify({"status": "error", "message": traceback.format_exc()}), 500


@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    data = request.get_json()
    try:
        doc = generate_docx(data)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        response = make_response(buf.read())
        response.headers["Content-Type"] = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response.headers["Content-Disposition"] = "attachment; filename=ta_schedule.docx"
        return response
    except Exception:
        traceback.print_exc()
        return jsonify({"error": traceback.format_exc()}), 500


# ── solver ───────────────────────────────────────────────────────────────────

def solve(data):
    try:
        from ortools.sat.python import cp_model
    except ImportError:
        return {"status": "error", "message": "ortools not installed. Run: pip install ortools"}

    model = cp_model.CpModel()

    roles_map = {r["id"]: r for r in data.get("roles", [])}
    labs = data.get("labs", [])
    tas = data.get("tas", [])
    assignments_in = data.get("assignments", [])
    gc_map = {gc["id"]: gc for gc in data.get("grad_courses", [])}

    if not labs or not tas:
        return {"status": "feasible", "assignments": assignments_in, "diagnostics": {}}

    locked_set = {
        (a["lab_id"], a["role_id"], a["ta_id"])
        for a in assignments_in
        if a.get("locked")
    }

    # ── decision variables ──
    x = {}
    real_vars = []
    for lab in labs:
        for rr in lab.get("roles", []):
            role_id = rr["role_id"]
            req_exp = rr.get("required_experienced", False)
            for ta in tas:
                key = (lab["id"], role_id, ta["id"])

                # experience hard-exclude
                if req_exp and ta.get("experience") != "experienced":
                    x[key] = model.NewConstant(0)
                    continue

                # availability hard-exclude
                unavail = False
                for gc_id in ta.get("grad_course_ids", []):
                    gc = gc_map.get(gc_id)
                    if gc and gc["day"] == lab.get("day") and times_overlap(
                        lab["start_min"], lab["end_min"], gc["start_min"], gc["end_min"]
                    ):
                        unavail = True
                        break
                if not unavail:
                    for oc in ta.get("other_commitments", []):
                        if oc["day"] == lab.get("day") and times_overlap(
                            lab["start_min"], lab["end_min"], oc["start_min"], oc["end_min"]
                        ):
                            unavail = True
                            break
                if unavail:
                    x[key] = model.NewConstant(0)
                    continue

                var = model.NewBoolVar(f"x_{lab['id']}_{role_id}_{ta['id']}")
                if key in locked_set:
                    model.Add(var == 1)
                else:
                    real_vars.append(var)
                x[key] = var

    # ── constraint 1: role count (soft upper bound, maximize fill) ──
    for lab in labs:
        for rr in lab.get("roles", []):
            role_id = rr["role_id"]
            count = rr.get("count", 1)
            ta_vars = [x[(lab["id"], role_id, ta["id"])]
                       for ta in tas if (lab["id"], role_id, ta["id"]) in x]
            model.Add(sum(ta_vars) <= count)

    if real_vars:
        model.Maximize(sum(real_vars))

    # ── constraint 2: SE cap ──
    SE_SCALE = 100
    for ta in tas:
        max_se = int(ta.get("max_se", 2.0) * SE_SCALE)
        outside_se = sum(int(d.get("se_value", 0) * SE_SCALE)
                         for d in ta.get("outside_duties", []))
        se_terms = []
        for lab in labs:
            for rr in lab.get("roles", []):
                key = (lab["id"], rr["role_id"], ta["id"])
                if key in x:
                    se_val = int(roles_map.get(rr["role_id"], {}).get("se_value", 1.0) * SE_SCALE)
                    se_terms.append(x[key] * se_val)
        if se_terms:
            model.Add(sum(se_terms) + outside_se <= max_se)

    # ── constraint 3: no double-booking ──
    for ta in tas:
        for i, lab1 in enumerate(labs):
            for lab2 in labs[i + 1:]:
                if lab1.get("day") != lab2.get("day"):
                    continue
                if not times_overlap(
                    lab1.get("start_min", 0), lab1.get("end_min", 0),
                    lab2.get("start_min", 0), lab2.get("end_min", 0),
                ):
                    continue
                v1 = [x[(lab1["id"], rr["role_id"], ta["id"])]
                      for rr in lab1.get("roles", [])
                      if (lab1["id"], rr["role_id"], ta["id"]) in x]
                v2 = [x[(lab2["id"], rr["role_id"], ta["id"])]
                      for rr in lab2.get("roles", [])
                      if (lab2["id"], rr["role_id"], ta["id"]) in x]
                if v1 and v2:
                    model.Add(sum(v1) + sum(v2) <= 1)

    solver = cp_model.CpSolver()
    solver.parameters.max_time_in_seconds = 30
    status = solver.Solve(model)

    if status in (cp_model.OPTIMAL, cp_model.FEASIBLE):
        new_assignments = [a for a in assignments_in if a.get("locked")]
        for lab in labs:
            for rr in lab.get("roles", []):
                role_id = rr["role_id"]
                for ta in tas:
                    key = (lab["id"], role_id, ta["id"])
                    if key in locked_set:
                        continue
                    if key in x and solver.Value(x[key]) == 1:
                        new_assignments.append({
                            "lab_id": lab["id"],
                            "role_id": role_id,
                            "ta_id": ta["id"],
                            "locked": False,
                        })

        # compute unfilled roles
        unfilled = []
        for lab in labs:
            for rr in lab.get("roles", []):
                role_id = rr["role_id"]
                count = rr.get("count", 1)
                assigned = sum(
                    1 for a in new_assignments
                    if a["lab_id"] == lab["id"] and a["role_id"] == role_id
                )
                if assigned < count:
                    role_label = roles_map.get(role_id, {}).get("label", role_id)
                    unfilled.append({
                        "lab_id": lab["id"],
                        "lab_name": lab["name"],
                        "role_id": role_id,
                        "role_label": role_label,
                        "assigned": assigned,
                        "required": count,
                    })

        solve_status = "partial" if unfilled else ("optimal" if status == cp_model.OPTIMAL else "feasible")
        return {
            "status": solve_status,
            "assignments": new_assignments,
            "diagnostics": {"unfilled_roles": unfilled},
        }
    else:
        return {
            "status": "infeasible",
            "assignments": [a for a in assignments_in if a.get("locked")],
            "diagnostics": {"unfilled_roles": [], "error": "Locked assignments conflict with constraints."},
        }


# ── DOCX export ──────────────────────────────────────────────────────────────

def generate_docx(data):
    from docx import Document

    doc = Document()
    doc.add_heading("TA Schedule", 0)

    roles_map = {r["id"]: r for r in data.get("roles", [])}
    tas_map = {t["id"]: t for t in data.get("tas", [])}
    labs_map = {l["id"]: l for l in data.get("labs", [])}
    assignments = data.get("assignments", [])
    day_long = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
    day_short = ["Mon", "Tue", "Wed", "Thu", "Fri"]

    # Lab-centric
    doc.add_heading("Lab Assignments", 1)
    for lab in data.get("labs", []):
        doc.add_heading(lab["name"], 2)
        day = lab.get("day", 0)
        doc.add_paragraph(
            f"{day_long[day]}, {fmt_time(lab.get('start_min', 480))} – {fmt_time(lab.get('end_min', 540))}"
        )
        lab_asgn = [a for a in assignments if a["lab_id"] == lab["id"]]
        if lab_asgn:
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Role", "TA", "Status"
            for a in lab_asgn:
                row = tbl.add_row().cells
                row[0].text = roles_map.get(a["role_id"], {}).get("label", a.get("role_id", ""))
                row[1].text = tas_map.get(a["ta_id"], {}).get("name", a.get("ta_id", ""))
                row[2].text = "Locked" if a.get("locked") else "Assigned"
        else:
            doc.add_paragraph("No assignments")
        doc.add_paragraph()

    # TA-centric
    doc.add_heading("TA Assignments", 1)
    for ta in data.get("tas", []):
        doc.add_heading(ta["name"], 2)
        doc.add_paragraph(
            f"{ta.get('experience','').capitalize()}, Max SE: {ta.get('max_se', 2.0)}"
        )
        ta_asgn = [a for a in assignments if a["ta_id"] == ta["id"]]
        if ta_asgn:
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Lab", "Role", "Time", "SE"
            total_se = 0.0
            for a in sorted(
                ta_asgn,
                key=lambda a: (
                    labs_map.get(a["lab_id"], {}).get("day", 0),
                    labs_map.get(a["lab_id"], {}).get("start_min", 0),
                ),
            ):
                lab = labs_map.get(a["lab_id"], {})
                role = roles_map.get(a["role_id"], {})
                se = role.get("se_value", 0)
                total_se += se
                row = tbl.add_row().cells
                row[0].text = lab.get("name", "")
                row[1].text = role.get("label", "")
                if lab:
                    d = lab.get("day", 0)
                    row[2].text = (
                        f"{day_short[d]} {fmt_time(lab['start_min'])}–{fmt_time(lab['end_min'])}"
                    )
                row[3].text = str(se)
            tot = tbl.add_row().cells
            tot[0].text = "Total SE"
            tot[3].text = f"{total_se} / {ta.get('max_se', 2.0)}"
        else:
            doc.add_paragraph("No assignments")
        doc.add_paragraph()

    return doc


if __name__ == "__main__":
    os.makedirs("static", exist_ok=True)
    app.run(debug=True, port=5050)
